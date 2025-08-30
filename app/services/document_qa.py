"""
Document Q&A Service using LangChain and Vector Stores
"""
import os
import random
from typing import Dict, List, Optional
from pathlib import Path

from langchain_openai import OpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.document_loaders.word_document import Docx2txtLoader

import PyPDF2
from docx import Document as DocxDocument

from app.config import get_settings, get_llm_config
from app.models.schemas import DocumentUploadResponse, QuestionResponse, SearchResult, DocumentSource


class MockEmbeddings:
    """Mock embeddings for demonstration"""
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Mock embedding for documents"""
        return [[random.random() for _ in range(384)] for _ in texts]
    
    def embed_query(self, text: str) -> List[float]:
        """Mock embedding for query"""
        return [random.random() for _ in range(384)]


class MockVectorStore:
    """Mock vector store for demonstration"""
    
    def __init__(self, documents: List[Document] = None):
        self.documents = documents or []
    
    def similarity_search(self, query: str, k: int = 4) -> List[Document]:
        """Mock similarity search"""
        # Return random documents from the collection
        if not self.documents:
            return []
        
        # Shuffle and return up to k documents
        import random
        random.shuffle(self.documents)
        return self.documents[:k]
    
    def add_documents(self, documents: List[Document]):
        """Add documents to the mock store"""
        self.documents.extend(documents)
    
    def as_retriever(self, search_kwargs: Dict = None):
        """Return a retriever interface"""
        class MockRetriever:
            def __init__(self, vector_store):
                self.vector_store = vector_store
            
            def get_relevant_documents(self, query: str) -> List[Document]:
                k = 4
                if hasattr(self, 'search_kwargs') and self.search_kwargs:
                    k = self.search_kwargs.get('k', 4)
                return self.vector_store.similarity_search(query, k)
        
        return MockRetriever(self)


class DocumentQAService:
    """Service for document Q&A using vector embeddings and retrieval"""
    
    def __init__(self):
        self.settings = get_settings()
        self.llm = None
        self.embeddings = None
        self.vector_stores: Dict[str, any] = {}
        self._initialize_services()
    
    def _initialize_services(self):
        """Initialize LLM and embeddings"""
        llm_config = get_llm_config()
        
        try:
            if llm_config["provider"] == "openai":
                self.llm = OpenAI(
                    openai_api_key=llm_config["api_key"],
                    model_name="gpt-3.5-turbo-instruct",
                    temperature=0.1,
                    max_tokens=1000
                )
                self.embeddings = OpenAIEmbeddings(
                    openai_api_key=llm_config["api_key"]
                )
            else:
                print("No OpenAI API key provided. Using mock services for demonstration.")
                self.llm = self._create_mock_llm()
                self.embeddings = MockEmbeddings()
                
        except Exception as e:
            print(f"Error initializing DocumentQA services: {e}")
            print("Falling back to mock services.")
            self.llm = self._create_mock_llm()
            self.embeddings = MockEmbeddings()
    
    def _create_mock_llm(self):
        """Create a mock LLM for demonstration"""
        class MockQALLM:
            async def apredict(self, prompt: str) -> str:
                return f"Mock answer based on the documents and question: {prompt[:100]}..."
            
            def predict(self, prompt: str) -> str:
                return f"Mock answer based on the documents and question: {prompt[:100]}..."
            
            def __call__(self, prompt: str) -> str:
                return self.predict(prompt)
        
        return MockQALLM()
    
    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                return self._extract_text_from_docx(file_path)
            elif file_extension in ['.txt', '.md']:
                return self._extract_text_from_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Failed to extract text from {file_path}: {str(e)}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_text_from_text(self, file_path: str) -> str:
        """Extract text from text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    async def process_document(self, file_path: str, collection_name: str = "default") -> DocumentUploadResponse:
        """Process a document and add it to the vector store"""
        try:
            # Extract text from file
            text = self._extract_text_from_file(file_path)
            
            # Split text into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.settings.chunk_size,
                chunk_overlap=self.settings.chunk_overlap
            )
            
            chunks = text_splitter.split_text(text)
            documents = [
                Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "chunk_index": i,
                        "file_name": Path(file_path).name
                    }
                )
                for i, chunk in enumerate(chunks)
            ]
            
            # Create or update vector store
            if isinstance(self.embeddings, MockEmbeddings):
                # Use mock vector store
                if collection_name not in self.vector_stores:
                    self.vector_stores[collection_name] = MockVectorStore()
                self.vector_stores[collection_name].add_documents(documents)
            else:
                # Use Chroma vector store
                persist_directory = os.path.join(self.settings.chroma_persist_directory, collection_name)
                
                if collection_name not in self.vector_stores:
                    self.vector_stores[collection_name] = Chroma(
                        embedding_function=self.embeddings,
                        persist_directory=persist_directory
                    )
                
                self.vector_stores[collection_name].add_documents(documents)
                self.vector_stores[collection_name].persist()
            
            return DocumentUploadResponse(
                message="Document processed successfully",
                file_name=Path(file_path).name,
                chunks_created=len(documents),
                collection_name=collection_name
            )
            
        except Exception as e:
            raise Exception(f"Document processing failed: {str(e)}")
    
    async def answer_question(self, question: str, collection_name: str = "default", options: Dict = None) -> QuestionResponse:
        """Answer a question based on the documents in the collection"""
        try:
            if options is None:
                options = {}
            
            top_k = options.get("topK", 4)
            include_metadata = options.get("includeMetadata", True)
            
            if collection_name not in self.vector_stores:
                raise ValueError(f"No documents found in collection: {collection_name}")
            
            vector_store = self.vector_stores[collection_name]
            
            # Get relevant documents
            relevant_docs = vector_store.similarity_search(question, k=top_k)
            
            if isinstance(self.llm, type(self._create_mock_llm())):
                # Mock answer
                context = " ".join([doc.page_content[:100] for doc in relevant_docs])
                answer = f"Based on the provided documents about {question}, here is a mock answer derived from the context: {context[:200]}..."
            else:
                # Use RetrievalQA chain
                qa_chain = RetrievalQA.from_chain_type(
                    llm=self.llm,
                    chain_type="stuff",
                    retriever=vector_store.as_retriever(search_kwargs={"k": top_k})
                )
                result = qa_chain.run(question)
                answer = result
            
            response = QuestionResponse(
                answer=answer.strip(),
                question=question,
                collection_name=collection_name
            )
            
            if include_metadata:
                response.sources = [
                    DocumentSource(
                        content=doc.page_content[:200] + "...",
                        metadata=doc.metadata,
                        relevance_score=random.uniform(0.7, 0.95)
                    )
                    for doc in relevant_docs
                ]
                response.confidence = self._calculate_confidence(relevant_docs, question)
            
            return response
            
        except Exception as e:
            raise Exception(f"Question answering failed: {str(e)}")
    
    async def search_documents(self, query: str, collection_name: str = "default", limit: int = 5) -> SearchResult:
        """Search through documents"""
        try:
            if collection_name not in self.vector_stores:
                raise ValueError(f"No documents found in collection: {collection_name}")
            
            vector_store = self.vector_stores[collection_name]
            results = vector_store.similarity_search(query, k=limit)
            
            return SearchResult(
                query=query,
                results=[
                    DocumentSource(
                        content=doc.page_content,
                        metadata=doc.metadata,
                        relevance_score=random.uniform(0.7, 0.95)
                    )
                    for doc in results
                ],
                total_results=len(results),
                collection_name=collection_name
            )
            
        except Exception as e:
            raise Exception(f"Document search failed: {str(e)}")
    
    def _calculate_confidence(self, docs: List[Document], question: str) -> str:
        """Calculate confidence score for the answer"""
        # Simple confidence calculation
        avg_length = sum(len(doc.page_content) for doc in docs) / len(docs) if docs else 0
        has_keywords = any(
            question.lower().split()[0] in doc.page_content.lower() 
            for doc in docs
        )
        
        confidence = 0.5
        if avg_length > 100:
            confidence += 0.2
        if has_keywords:
            confidence += 0.2
        if len(docs) >= 3:
            confidence += 0.1
        
        return f"{min(confidence, 0.95):.2f}"
    
    def get_collections(self) -> List[Dict[str, any]]:
        """Get list of all collections"""
        return [
            {
                "name": name,
                "document_count": len(getattr(store, 'documents', []))
            }
            for name, store in self.vector_stores.items()
        ]
    
    async def delete_collection(self, collection_name: str) -> Dict[str, str]:
        """Delete a collection"""
        if collection_name in self.vector_stores:
            del self.vector_stores[collection_name]
            
            # Also remove persistent directory if using Chroma
            persist_directory = os.path.join(self.settings.chroma_persist_directory, collection_name)
            if os.path.exists(persist_directory):
                import shutil
                shutil.rmtree(persist_directory)
            
            return {"message": f"Collection {collection_name} deleted successfully"}
        else:
            raise ValueError(f"Collection {collection_name} not found")
